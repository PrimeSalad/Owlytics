import sys; sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

SRC  = r"C:\Users\gland\Videos\student monitoring system\docs\ACCOMPLISHMENT-REPORT-2025-2026-format template.docx"
DEST = r"C:\Users\gland\Videos\student monitoring system\server\src\assets\template.docx"

doc = Document(SRC)

def make_tag_row(text, source_row):
    tr = deepcopy(source_row._tr)
    for tc in tr.findall(qn('w:tc')):
        for p_el in tc.findall(qn('w:p')):
            for r_el in p_el.findall(qn('w:r')):
                p_el.remove(r_el)
    first_tc = tr.findall(qn('w:tc'))[0]
    first_p  = first_tc.findall(qn('w:p'))[0]
    r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = text; r.append(t); first_p.append(r)
    return tr

# Summary table (T0): wrap data row with {#activities}
t0 = doc.tables[0]
data_tr = t0.rows[1]._tr
t0._tbl.remove(data_tr)
t0._tbl.append(make_tag_row('{#activities}', t0.rows[0]))
t0._tbl.append(data_tr)
t0._tbl.append(make_tag_row('{/activities}', t0.rows[0]))
print('OK T0 loop added')

# Documentation table (T1): wrap all rows with {#activities}
t1 = doc.tables[1]
rows_tr = [deepcopy(r._tr) for r in t1.rows]
for tr in list(t1._tbl.findall(qn('w:tr'))):
    t1._tbl.remove(tr)
t1._tbl.append(make_tag_row('{#activities}', t0.rows[0]))
for tr in rows_tr:
    t1._tbl.append(tr)
t1._tbl.append(make_tag_row('{/activities}', t0.rows[0]))
print('OK T1 loop added')

doc.save(DEST)
print(f'OK Saved {DEST}')
